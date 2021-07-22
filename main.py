#!/usr/bin/python
# -*- coding: UTF-8 -*-
 
import sys
import argparse
import docx
import json
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import OxmlElement, ns
from docx.shared import Cm, Pt

# Инициализация входных параметров.
def createParser ():
    parser = argparse.ArgumentParser()
    parser.add_argument ('inputfile', type=str)
    parser.add_argument ('data', type=str)
    parser.add_argument ('otputfile', type=str)
    return parser

def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

# Тонкая работа с XML для вывода особых штук - номера страницы.
def append_special_thing(paragraph, thing):
    #--- Добавляем номер страницы
    # запускаем динамическое обновление параграфа
    page_num_run = paragraph.add_run()
    # обозначаем начало позиции вывода
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    # задаем вывод текущего значения страницы PAGE (всего страниц NUMPAGES)
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = thing
    # обозначаем конец позиции вывода
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    # добавляем все в наш параграф (который формируется динамически)
    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

# Вывести номер страницы в нужном виде.
def add_page_number(paragraph):
    # выравниваем параграф по правому краю
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    # Добавляем текст "Страница"
    paragraph.add_run("Страница ")
    # Добавляем номер страницы
    append_special_thing(paragraph, "PAGE")
    # Добавляем текст "из"
    paragraph.add_run(" из ")
    # Добавляем количество страниц
    append_special_thing(paragraph, "NUMPAGES")

# Добавить строку с договором.    
def add_contract(paragraph):
    paragraph.add_run("Договор № ______")
    run = paragraph.add_run()
    run.add_tab()

# Установить абзацу правильные стиль и формат. Много хардкода.
def make_beauty(paragraph):
    paragraph.style.font.size = Pt(7)
    paragraph.style.font.name = 'Verdana'
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(10.2)

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

# Основное действие.
if __name__ == '__main__':
    # Парсер параметров.
    parser = createParser()
    arguments = parser.parse_args(sys.argv[1:])

    # Чтение документа-исходника.
    doc = docx.Document(arguments.inputfile)

    # Чтение входящих параметров.
    data = json.load(open(arguments.data, encoding='utf-8'))
    
    section = doc.sections[0]
    footer = section.footer

    # Чистим все абзацы футера, на всякий случай.
    # footer.paragraphs._r.RemoveAll()
    for i in footer.paragraphs:
       delete_paragraph(i) 

    # Получаем ширину страницы в универсальных единицах. Это потом пригодится.
    page_width = section.page_width

    # Не смотря на то, что параграфы очистили, первый элемент все равно есть. Не знаю почему.
    paragraph = footer.add_paragraph()

    # Добавим немного стиля.
    make_beauty(paragraph)

    # Добавляем таб-стоп, чтобы разнести элементы по разным сторонам строки.
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stop = tab_stops.add_tab_stop(page_width, WD_TAB_ALIGNMENT.RIGHT)

    # Добавим строку с договором.
    add_contract(paragraph)
    # Добавим номер страницы.
    add_page_number(paragraph)

    # Добавим абзац с визами.
    par = footer.add_paragraph(data.get("text"))
    
    # Добавим немного стиля.
    make_beauty(par)

    # Сохраним документ.
    doc.save(arguments.otputfile)